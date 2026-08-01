import importlib.util
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn

from core.security import tool_execution_policy
from core.document_text import extract_document_text


ROOT = Path(__file__).resolve().parents[2]
ALIX = ROOT / "starter_packs" / "document_heavy" / "Alix-AI"


def _load(name, path):
    spec = importlib.util.spec_from_file_location(name, path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def test_bundled_intake_forms_are_valid_protected_content_control_documents():
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    from lxml import etree

    for template in sorted((ALIX / "templates").glob("intake_*/template.docx")):
        with ZipFile(template) as archive:
            assert archive.testzip() is None
            assert not any("vbaProject" in name for name in archive.namelist())
            for name in archive.namelist():
                if name.endswith(".rels"):
                    assert b'TargetMode="External"' not in archive.read(name)
            document_xml = etree.fromstring(archive.read("word/document.xml"))
            settings_xml = etree.fromstring(archive.read("word/settings.xml"))
        tags = document_xml.xpath(".//w:sdt/w:sdtPr/w:tag/@w:val", namespaces=namespace)
        protection = settings_xml.xpath(".//w:documentProtection", namespaces=namespace)
        assert len(tags) >= 10
        assert len(tags) == len(set(tags))
        assert protection
        assert protection[0].get(f"{{{namespace['w']}}}edit") == "forms"
        assert protection[0].get(f"{{{namespace['w']}}}enforcement") == "1"


def test_fillable_form_round_trip_is_idempotent_and_keeps_unrelated_lines(tmp_path, monkeypatch):
    engine_module = _load("fillable_document_engine", ALIX / "business" / "document_engine.py")
    reader_module = _load("fillable_form_reader", ALIX / "tools" / "read_filled_form.py")
    path = tmp_path / "intake.docx"
    document = Document()
    document.add_paragraph("1. Full legal name")
    document.add_paragraph("____________________________")
    document.add_paragraph("Instructions for the next section")
    document.add_paragraph("____________________________")
    document.save(path)

    engine = engine_module.DocumentEngine(str(path))
    assert engine.make_fillable(str(path)) == 1
    engine.apply_forms_protection(str(path))
    assert engine.make_fillable(str(path)) == 0

    filled = Document(path)
    controls = list(filled.element.body.iter(qn("w:sdt")))
    assert len(controls) == 1
    assert controls[0].find(f'.//{qn("w:tag")}').get(qn("w:val")) == "q1_1"
    text = controls[0].find(f'.//{qn("w:sdtContent")}/{qn("w:r")}/{qn("w:t")}')
    text.text = "Synthetic Client"
    filled.save(path)

    extraction = extract_document_text(str(path))
    assert "Synthetic Client" in extraction.text

    monkeypatch.setattr(reader_module, "require_allowed_path", lambda value: str(Path(value).resolve()))
    result = reader_module.execute(str(path))
    assert '"q1_1"' in result
    assert '"answer": "Synthetic Client"' in result
    assert "1 answered, 0 left blank" in result


def test_fillable_form_tools_enforce_beta_security_boundaries():
    reader_module = _load("fillable_form_reader_security", ALIX / "tools" / "read_filled_form.py")
    assert reader_module.execute("/etc/passwd").startswith("Error: Path is not inside")
    allowed, reason = tool_execution_policy("build_fillable_form", {"template_name": "intake_probate"})
    assert allowed is False
    assert "developer mode" in reason
