import os
import csv
import yaml
import sys
import subprocess

# Optional imports handled gracefully
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None

TOOL_DEFINITION = {
    "name": "read_document",
    "description": "Reads and extracts plain text content from any supported document format in the inbox or workspace (PDF, DOCX, DOC, TXT, CSV, XLSX, RTF, images).",
    "parameters": {
        "type": "object",
        "properties": {
            "file_name": {
                "type": "string",
                "description": "The name of the file to read (must be in workspace/inbox, templates, or output folder, or an absolute path)."
            },
            "directory_type": {
                "type": "string",
                "enum": ["inbox", "templates", "output", "skills"],
                "description": "The folder where the file resides. Defaults to inbox.",
                "default": "inbox"
            }
        },
        "required": ["file_name"]
    }
}

def get_config():
    tools_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(tools_dir)
    config_path = os.path.join(project_dir, "config.yaml")
    if os.path.exists(config_path):
        with open(config_path, "r") as f:
            try:
                return yaml.safe_load(f)
            except Exception:
                pass
    return {}

def resolve_path(file_name, directory_type):
    if os.path.isabs(file_name):
        return file_name
        
    config = get_config()
    paths = config.get("paths", {})
    fallback_paths = {
        "inbox": "./workspace/inbox",
        "output": "./workspace/output",
        "templates": "./templates",
        "skills": "./skills"
    }
    dir_path = paths.get(directory_type, fallback_paths.get(directory_type, "./workspace/inbox"))
    return os.path.join(dir_path, file_name)

def read_pdf(path):
    if not fitz:
        return "Error: PyMuPDF (pymupdf) is not installed. Run 'pip install pymupdf'."
    
    try:
        doc = fitz.open(path)
        text_content = []
        for i, page in enumerate(doc):
            text_content.append(f"--- PAGE {i+1} ---")
            text_content.append(page.get_text())
        return "\n".join(text_content)
    except Exception as e:
        return f"Error reading PDF: {e}"

def read_docx(path):
    if not docx:
        return "Error: python-docx is not installed. Run 'pip install python-docx'."
    
    try:
        doc_obj = docx.Document(path)
        text_content = []
        for para in doc_obj.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
                
        # Also extract tables if present
        for i, table in enumerate(doc_obj.tables):
            text_content.append(f"\n--- TABLE {i+1} ---")
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                text_content.append(" | ".join(row_data))
                
        return "\n".join(text_content)
    except Exception as e:
        return f"Error reading DOCX: {e}"

def read_doc_legacy(path):
    """Fallback reader for legacy .doc files using antiword if available, or basic ASCII dump."""
    # Check if antiword is installed
    import shutil
    if shutil.which("antiword"):
        try:
            res = subprocess.run(["antiword", path], capture_output=True, text=True, check=True)
            return res.stdout
        except Exception as e:
            return f"Error calling antiword: {e}"
            
    # Pure-python basic ASCII/printable extraction fallback for .doc
    try:
        with open(path, "rb") as f:
            data = f.read()
        # Find printable ascii sequences (a naive attempt to get text from binary stream)
        import string
        printable = set(string.printable.encode('ascii'))
        text_chars = []
        for b in data:
            if b in printable:
                text_chars.append(chr(b))
            elif b == 0:
                text_chars.append(' ')
        text = "".join(text_chars)
        # Clean up excessive spaces/newlines
        import re
        text = re.sub(r'\s+', ' ', text).strip()
        return f"[WARNING: Naive text extraction used for legacy .doc file. antiword not installed.]\n\nExtracted Text:\n{text[:2000]}..."
    except Exception as e:
        return f"Error reading legacy DOC: {e}"

def read_xlsx(path):
    if not openpyxl:
        return "Error: openpyxl is not installed. Run 'pip install openpyxl'."
        
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        content = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            content.append(f"--- SHEET: {sheet_name} ---")
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    content.append(row_str)
        return "\n".join(content)
    except Exception as e:
        return f"Error reading XLSX: {e}"

def read_csv(path):
    try:
        content = []
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                content.append(" , ".join(row))
        return "\n".join(content)
    except Exception as e:
        return f"Error reading CSV: {e}"

def read_rtf(path):
    if not rtf_to_text:
        return "Error: striprtf is not installed. Run 'pip install striprtf'."
        
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            rtf_content = f.read()
        return rtf_to_text(rtf_content)
    except Exception as e:
        return f"Error reading RTF: {e}"

def read_image_ocr(path):
    """Tries using pytesseract if available, else returns file metadata."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        return f"--- OCR RESULTS ({os.path.basename(path)}) ---\n{text}"
    except ImportError:
        return f"[IMAGE FILE: {os.path.basename(path)}] pytesseract or PIL is not installed. Install pytesseract to extract text from images."
    except Exception as e:
        return f"Error executing OCR: {e}"

def execute(file_name, directory_type="inbox"):
    full_path = resolve_path(file_name, directory_type)
    
    if not os.path.exists(full_path):
        return f"Error: File '{file_name}' not found at path: {full_path}"
        
    ext = os.path.splitext(file_name)[1].lower()
    
    if ext == ".pdf":
        return read_pdf(full_path)
    elif ext == ".docx":
        return read_docx(full_path)
    elif ext == ".doc":
        return read_doc_legacy(full_path)
    elif ext in [".xlsx", ".xls"]:
        return read_xlsx(full_path)
    elif ext == ".csv":
        return read_csv(full_path)
    elif ext == ".rtf":
        return read_rtf(full_path)
    elif ext in [".txt", ".md", ".log", ".yaml", ".yml", ".json"]:
        try:
            with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            return f"Error reading text file: {e}"
    elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
        return read_image_ocr(full_path)
    else:
        # Try reading as plain text anyway
        try:
            with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return f"[WARNING: Unknown extension '{ext}'. Read as text.]\n\n{content}"
        except Exception:
            return f"Error: Unsupported file format '{ext}'."
