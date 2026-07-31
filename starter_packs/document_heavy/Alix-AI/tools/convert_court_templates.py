import os

def _find_aimaos_root():
    p = os.path.dirname(os.path.abspath(__file__))
    while p != os.path.dirname(p) and not os.path.exists(os.path.join(p, "aimaos_config.yaml")):
        p = os.path.dirname(p)
    return p
AIMAOS_ROOT = os.environ.get("AIMAOS_ROOT") or _find_aimaos_root()
import glob
import re
import yaml
import docx

DOWNLOADS_DIR = os.path.expanduser("~/Downloads")
TEMPLATES_DIR = os.path.join(AIMAOS_ROOT, "Alix-AI/templates")

FORM_MAPPING = {
    "12.900(h) ADA.docx": {
        "folder": "form_12_900_h",
        "name": "Notice of Related Cases (Form 12.900h)",
        "description": "Florida Family Law Rules of Procedure Form 12.900(h) - Notice of Related Cases",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number", "division"]
    },
    "12.901(a) ADA.docx": {
        "folder": "form_12_901_a",
        "name": "Joint Petition for Simplified Dissolution of Marriage (Form 12.901a)",
        "description": "Florida Family Law Rules of Procedure Form 12.901(a) - Joint Petition for Simplified Dissolution of Marriage",
        "fields": ["husband_name", "wife_name", "client_name", "spouse_name", "county", "circuit_number", "case_number", "division"]
    },
    "12.902(c) ADA.docx": {
        "folder": "form_12_902_c",
        "name": "Family Law Financial Affidavit Long Form (Form 12.902c)",
        "description": "Florida Family Law Rules of Procedure Form 12.902(c) - Family Law Financial Affidavit (Long Form)",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number", "division", "occupation", "employer"]
    },
    "12.902(d) ADA.docx": {
        "folder": "form_12_902_d",
        "name": "Uniform Child Custody Jurisdiction Affidavit (Form 12.902d)",
        "description": "Florida Family Law Rules of Procedure Form 12.902(d) - Uniform Child Custody Jurisdiction and Enforcement Act (UCCJEA) Affidavit",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number", "children_names"]
    },
    "12.902(e) ADA.docx": {
        "folder": "form_12_902_e",
        "name": "Child Support Guidelines Worksheet (Form 12.902e)",
        "description": "Florida Family Law Rules of Procedure Form 12.902(e) - Child Support Guidelines Worksheet",
        "fields": ["father_name", "mother_name", "client_name", "county", "circuit_number", "case_number", "num_children"]
    },
    "12.902(f)(3) ADA.docx": {
        "folder": "form_12_902_f3",
        "name": "Marital Settlement Agreement (Form 12.902f3)",
        "description": "Florida Family Law Rules of Procedure Form 12.902(f)(3) - Marital Settlement Agreement for Simplified Dissolution of Marriage",
        "fields": ["husband_name", "wife_name", "client_name", "spouse_name", "county", "circuit_number", "case_number"]
    },
    "12.902(j) ADA.docx": {
        "folder": "form_12_902_j",
        "name": "Notice of Social Security Number (Form 12.902j)",
        "description": "Florida Family Law Rules of Procedure Form 12.902(j) - Notice of Social Security Number",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number"]
    },
    "12.902(k) ADA.docx": {
        "folder": "form_12_902_k",
        "name": "Notice of Joint Filing (Form 12.902k)",
        "description": "Florida Family Law Rules of Procedure Form 12.902(k) - Notice of Joint Filing",
        "fields": ["client_name", "spouse_name", "county", "circuit_number", "case_number"]
    },
    "12.905(a) ADA.docx": {
        "folder": "form_12_905_a",
        "name": "Supplemental Petition to Modify Parenting Plan (Form 12.905a)",
        "description": "Florida Family Law Rules of Procedure Form 12.905(a) - Supplemental Petition to Modify Parenting Plan/Time-Sharing Schedule",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number"]
    },
    "12.905(b) ADA.docx": {
        "folder": "form_12_905_b",
        "name": "Supplemental Petition to Modify Child Support (Form 12.905b)",
        "description": "Florida Family Law Rules of Procedure Form 12.905(b) - Supplemental Petition for Modification of Child Support",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number"]
    },
    "12.913(b) ADA.docx": {
        "folder": "form_12_913_b",
        "name": "Affidavit of Diligent Search and Inquiry (Form 12.913b)",
        "description": "Florida Family Law Rules of Procedure Form 12.913(b) - Affidavit of Diligent Search and Inquiry",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number"]
    },
    "12.913(c) ADA.docx": {
        "folder": "form_12_913_c",
        "name": "Notice of Action for Dissolution of Marriage (Form 12.913c)",
        "description": "Florida Family Law Rules of Procedure Form 12.913(c) - Notice of Action for Dissolution of Marriage",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number"]
    },
    "12.914 ADA.docx": {
        "folder": "form_12_914",
        "name": "Certificate of Service (Form 12.914)",
        "description": "Florida Family Law Rules of Procedure Form 12.914 - Certificate of Service",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number"]
    },
    "12.915 ADA.docx": {
        "folder": "form_12_915",
        "name": "Designation of Current Mailing and E-Mail Address (Form 12.915)",
        "description": "Florida Family Law Rules of Procedure Form 12.915 - Designation of Current Mailing and E-Mail Address",
        "fields": ["client_name", "petitioner_name", "respondent_name", "client_address", "client_phone", "client_email"]
    },
    "12.923 ADA.docx": {
        "folder": "form_12_923",
        "name": "Notice of Hearing (Form 12.923)",
        "description": "Florida Family Law Rules of Procedure Form 12.923 - Notice of Hearing (General Form)",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number", "hearing_date", "hearing_time"]
    },
    "12.932 ADA.docx": {
        "folder": "form_12_932",
        "name": "Notice of Filing Social Security Number (Form 12.932)",
        "description": "Florida Family Law Rules of Procedure Form 12.932 - Notice of Filing Social Security Number",
        "fields": ["client_name", "petitioner_name", "respondent_name", "county", "circuit_number", "case_number"]
    },
    "12.982(a) ADA.docx": {
        "folder": "form_12_982_a",
        "name": "Petition for Change of Name Adult (Form 12.982a)",
        "description": "Florida Supreme Court Approved Family Law Form 12.982(a) - Petition for Change of Name (Adult)",
        "fields": ["client_name", "county", "circuit_number", "case_number", "client_address", "client_phone", "date_of_birth"]
    },
    "12.982(b) ADA.docx": {
        "folder": "form_12_982_b",
        "name": "Final Judgment of Change of Name Adult (Form 12.982b)",
        "description": "Florida Supreme Court Approved Family Law Form 12.982(b) - Final Judgment of Change of Name (Adult)",
        "fields": ["client_name", "new_name", "county", "circuit_number", "case_number"]
    },
    "12.982(f) ADA.docx": {
        "folder": "form_12_982_f",
        "name": "Petition for Change of Name Minor Child (Form 12.982f)",
        "description": "Florida Supreme Court Approved Family Law Form 12.982(f) - Petition for Change of Name (Minor Child)",
        "fields": ["parent_name", "child_name", "county", "circuit_number", "case_number", "child_dob"]
    },
    "12.982(g) ADA.docx": {
        "folder": "form_12_982_g",
        "name": "Consent for Change of Name of Minor Child (Form 12.982g)",
        "description": "Florida Supreme Court Approved Family Law Form 12.982(g) - Consent for Change of Name of Minor Child",
        "fields": ["parent_name", "child_name", "county", "circuit_number", "case_number"]
    }
}

def inject_jinja2_tags(doc):
    """Replaces blank underlines and instructional placeholders with Jinja2 template tags."""
    
    # 1. Header & Caption placeholders
    for p in doc.paragraphs[:15]:
        text = p.text
        if "JUDICIAL CIRCUIT" in text:
            p.text = re.sub(r"(?:_{3,}|\s{5,})", " {{ circuit_number | default('[CIRCUIT NUMBER]') }} ", p.text)
        if "COUNTY, FLORIDA" in text:
            p.text = re.sub(r"(?:_{3,}|\s{5,})", " {{ county | default('[COUNTY]') }} ", p.text)
        if "Case No.:" in text:
            p.text = re.sub(r"Case No\.:\s*(?:_{3,}|\s{3,}|\t)", "Case No.: {{ case_number | default('[CASE NUMBER]') }}", p.text)
        if "Division:" in text:
            p.text = re.sub(r"Division:\s*(?:_{3,}|\s{3,}|\t)", "Division: {{ division | default('[DIVISION]') }}", p.text)

    # Caption party name underlines (e.g. ___________________, Petitioner / Respondent)
    for p in doc.paragraphs[:25]:
        if ("Petitioner" in p.text or "Respondent" in p.text) and ("____" in p.text or len(p.text.strip()) > 15):
            p.text = re.sub(r"_{3,}|\s{5,}", " {{ client_name | default(petitioner_name) }} ", p.text)

    # 2. Body Text placeholders
    for p in doc.paragraphs:
        text = p.text

        if "My complete present name is:" in text:
            p.text = re.sub(r"_{3,}", " {{ client_name }}", p.text)

        if "I request that my name be changed to:" in text:
            p.text = re.sub(r"_{3,}", " {{ new_name }}", p.text)

        if "I, {full legal name}" in text:
            p.text = re.sub(r"\{full legal name\}\s*_{3,}", "{{ client_name }}", p.text)

        if "I live in" in text and "County, Florida" in text:
            p.text = re.sub(r"_{3,}\s*County", "{{ county }} County", p.text)
            p.text = re.sub(r"\{street address\}\s*_{3,}", "{{ client_address }}", p.text)

        if "I was born on" in text:
            p.text = re.sub(r"\{date\}\s*_{3,}", "{{ date_of_birth }}", p.text)

        if "ORDERED that Petitioner’s present name," in text:
            p.text = re.sub(r"_{3,}", " {{ client_name }}", p.text)

        if "is changed to" in text and "by which Petitioner" in text:
            p.text = re.sub(r"_{3,}", " {{ new_name }}", p.text)

        if "Printed Name:" in text:
            p.text = re.sub(r"_{3,}", " {{ client_name }}", p.text)

        if "Address:" in text and "City" not in text:
            p.text = re.sub(r"_{3,}", " {{ client_address }}", p.text)

        if "City, State, Zip:" in text:
            p.text = re.sub(r"_{3,}", " {{ client_city_state_zip | default(client_address) }}", p.text)

        if "Telephone Number:" in text:
            p.text = re.sub(r"_{3,}", " {{ client_phone }}", p.text)

        if "Designated E-Mail Address(es):" in text:
            p.text = re.sub(r"_{3,}", " {{ client_email }}", p.text)

        if "COUNTY OF" in text and "STATE OF FLORIDA" not in text:
            p.text = re.sub(r"_{3,}|\t+", " {{ county }}", p.text)

        if "resident of" in text and "County, Florida" in text:
            p.text = re.sub(r"\t+|\s{3,}", " {{ county }} ", p.text)

        if "This cause came before the Court on" in text:
            p.text = re.sub(r"\{date\}\s*\t+", " {{ date | default('today') }}", p.text)

        if "DONE and ORDERED ON" in text:
            p.text = re.sub(r"_{3,}", " {{ date | default('today') }}", p.text, count=1)
            p.text = re.sub(r"_{3,}", " {{ county }}", p.text, count=1)

    # Standalone line underlines in captions or signatures
    for p in doc.paragraphs:
        stripped = p.text.strip()
        if stripped.startswith("____") and len(stripped.replace("_", "")) == 0:
            p.text = "{{ client_name }}"

def clean_and_convert():
    converted_count = 0
    for filename, meta in FORM_MAPPING.items():
        src_path = os.path.join(DOWNLOADS_DIR, filename)
        if not os.path.exists(src_path):
            print(f"Skipping missing file: {filename}")
            continue

        target_folder = os.path.join(TEMPLATES_DIR, meta["folder"])
        os.makedirs(target_folder, exist_ok=True)
        target_docx = os.path.join(target_folder, "template.docx")

        doc = docx.Document(src_path)

        # 1. Locate start of form (IN THE CIRCUIT COURT)
        start_idx = None
        for idx, p in enumerate(doc.paragraphs):
            text = p.text.upper().strip()
            if "IN THE CIRCUIT COURT" in text:
                start_idx = idx
                break

        # 2. Remove instructions preceding the form
        if start_idx is not None and start_idx > 0:
            for p in list(doc.paragraphs[:start_idx]):
                p._element.getparent().remove(p._element)

        # 3. Clear instruction headers and footers across all sections
        for s in doc.sections:
            for p in s.footer.paragraphs:
                p.text = ""
            for p in s.header.paragraphs:
                p.text = ""

        # 4. Inject Jinja2 Tags for full template auto-filling
        inject_jinja2_tags(doc)

        doc.save(target_docx)

        # 5. Create template.yaml metadata definition
        yaml_path = os.path.join(target_folder, "template.yaml")
        yaml_content = {
            "name": meta["name"],
            "description": meta["description"],
            "fields": {f: f"Value for {f.replace('_', ' ')}" for f in meta["fields"]},
            "default_format": "docx"
        }
        with open(yaml_path, "w") as f:
            yaml.dump(yaml_content, f, sort_keys=False)

        converted_count += 1

    print(f"\nCompleted re-converting {converted_count} court templates with Jinja2 auto-fill tags!")

if __name__ == "__main__":
    clean_and_convert()
