import os
import re
import subprocess
import json
import logging
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Text patterns that mean a "successfully rendered" document is still not
# actually finished -- a real, well-known failure class, not a hypothetical:
# - {{ tag }}: a Jinja tag Word split across separate XML runs (e.g. from
#   autocorrect/formatting) that docxtpl never recognized, so the field
#   silently never filled even though render() reported no error.
# - <TODO>/TBD/TODO: a placeholder note left by whoever authored/edited the
#   template, never meant to reach a client or court.
# - lorem ipsum / xxxx: boilerplate or filler text accidentally left in from
#   drafting the template itself.
# Anything matching after render means the document is broken, not just
# "rendered" -- flag it rather than reporting a plain success.
_LEAK_PATTERNS = [
    re.compile(r"\{\{.*?\}\}"),
    re.compile(r"<TODO>", re.IGNORECASE),
    re.compile(r"\bTBD\b"),
    re.compile(r"\blorem ipsum\b", re.IGNORECASE),
    re.compile(r"x{4,}", re.IGNORECASE),
]

# A numbered intake question ("12. Some question:") followed by one or more
# bare-underscore answer lines is the exact anti-pattern make_fillable()
# converts into real content controls -- see its docstring.
_QUESTION_NUMBER_RE = re.compile(r"^\s*(\d+)\.\s*(.*)$")
_UNDERSCORE_LINE_RE = re.compile(r"^_{10,}$")

class DocumentEngine:
    """
    Advanced Document Engine for Alix-AI.
    Renders Jinja2 templates inside Word .docx files, handles context validation,
    adds dynamic Table of Contents (TOC), and optionally converts output to PDF.
    """
    def __init__(self, template_path):
        self.template_path = os.path.abspath(template_path)
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template document not found: {self.template_path}")

    def _add_toc(self, paragraph):
        """Injects Word TOC field XML elements into a paragraph."""
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    def _force_update_toc(self, doc):
        """Forces Word to update fields (including TOC) upon next document open."""
        element = doc.settings.element.find(qn('w:updateFields'))
        if element is None:
            element = OxmlElement('w:updateFields')
            element.set(qn('w:val'), 'true')
            doc.settings.element.insert(0, element)

    def _build_sdt_field(self, tag, alias, field_id):
        """A single empty, plain-text OOXML content control (w:sdt).
        sdtContent is deliberately EMPTY, not a literal placeholder string --
        w:showingPlcHdr gives Word's grey hint-text display without that text
        actually being stored content, so "unanswered" is unambiguously an
        empty string on read-back (read_filled_form.py) rather than a magic
        sentinel a client could accidentally leave behind by tabbing past
        the field."""
        sdt = OxmlElement('w:sdt')
        sdtPr = OxmlElement('w:sdtPr')

        alias_el = OxmlElement('w:alias')
        alias_el.set(qn('w:val'), alias[:250])
        sdtPr.append(alias_el)

        tag_el = OxmlElement('w:tag')
        tag_el.set(qn('w:val'), tag)
        sdtPr.append(tag_el)

        id_el = OxmlElement('w:id')
        id_el.set(qn('w:val'), str(field_id))
        sdtPr.append(id_el)

        sdtPr.append(OxmlElement('w:showingPlcHdr'))
        sdtPr.append(OxmlElement('w:text'))
        sdt.append(sdtPr)
        sdt.append(OxmlElement('w:sdtEndPr'))

        sdt_content = OxmlElement('w:sdtContent')
        run = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = ''
        run.append(t)
        sdt_content.append(run)
        sdt.append(sdt_content)
        return sdt

    def make_fillable(self, output_path):
        """Converts the underscore-line answer-area anti-pattern (a bare
        '____...' paragraph directly under a numbered question) into real,
        empty content controls tagged by question number, e.g. 'q7_1' --
        stable and self-describing, no per-template hand-authored field map
        needed. A question with multiple consecutive answer lines (e.g. "list
        each child's name and DOB") gets 'q7_1', 'q7_2', ... so each line is
        independently readable. Everything else in the document (styles,
        headers, the Jinja {{ }} header fields populate_template still fills
        normally) is left untouched. Idempotent: a paragraph that no longer
        looks like a bare underscore line (already converted) is left alone.
        Returns the number of fields created."""
        doc = Document(output_path)
        used_ids = set()
        for id_element in doc.element.body.iter(qn('w:id')):
            try:
                used_ids.add(int(id_element.get(qn('w:val'))))
            except (TypeError, ValueError):
                continue

        def next_id():
            n = 1
            while n in used_ids:
                n += 1
            used_ids.add(n)
            return n

        current_num, current_text, line_index = None, None, 0
        fields_created = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            q_match = _QUESTION_NUMBER_RE.match(text)
            if q_match:
                current_num, current_text = q_match.group(1), q_match.group(2)
                line_index = 0
                continue
            if current_num and _UNDERSCORE_LINE_RE.match(text):
                line_index += 1
                tag = f"q{current_num}_{line_index}"
                alias = f"Q{current_num}: {current_text}" if line_index == 1 \
                    else f"Q{current_num} (line {line_index}): {current_text}"
                sdt = self._build_sdt_field(tag, alias, next_id())
                for run in list(para.runs):
                    run._element.getparent().remove(run._element)
                para._p.append(sdt)
                fields_created += 1
                continue
            if text:
                current_num, current_text, line_index = None, None, 0

        doc.save(output_path)
        return fields_created

    def apply_forms_protection(self, output_path):
        """Locks all non-field text from editing in Word (w:edit="forms")
        while leaving w:sdt content controls themselves editable -- the same
        pairing Word's own Restrict Editing -> "Filling in forms" uses.
        No hash/password: a legitimate "declared but unenforced" protection
        (Stop Protection needs no password), which is all that's needed
        here -- the goal is keeping a client from accidentally editing the
        printed questions, not securing the file against a determined user."""
        doc = Document(output_path)
        settings = doc.settings.element
        existing = settings.find(qn('w:documentProtection'))
        if existing is not None:
            settings.remove(existing)
        protection = OxmlElement('w:documentProtection')
        protection.set(qn('w:edit'), 'forms')
        protection.set(qn('w:enforcement'), '1')
        settings.insert(0, protection)
        doc.save(output_path)

    def validate_context(self, context, required_fields=None):
        """
        Validates context data dictionary against required fields.
        Fills missing fields with formatted placeholders.
        Returns (validated_context, missing_fields) -- the caller needs to
        know WHICH fields were left as placeholders, not just get a document
        that quietly contains them.
        """
        validated = dict(context or {})
        missing = []
        if required_fields:
            for field in required_fields:
                if field not in validated or validated[field] is None or validated[field] == "":
                    missing.append(field)
                    # Human-friendly missing placeholder
                    display_name = field.replace('_', ' ').title()
                    validated[field] = f"[{display_name} Required]"
        return validated, missing

    def _check_rendered_output(self, output_path):
        """Deterministic post-render sanity check -- catches a document that
        rendered without error but is still actually broken, rather than
        trusting that "no exception" means "correct." Returns a dict with
        `structural_error` (None if the file reopens cleanly) and
        `leak_tokens` (any leftover template tag or placeholder boilerplate
        still present -- see _LEAK_PATTERNS -- real failure signatures, not
        false positives)."""
        issues = {"structural_error": None, "leak_tokens": [], "unrendered_tags": []}
        try:
            doc = Document(output_path)
        except Exception as e:
            issues["structural_error"] = str(e)
            return issues

        texts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.extend(p.text for p in cell.paragraphs)
        for section in doc.sections:
            for part in (section.header, section.footer):
                texts.extend(p.text for p in part.paragraphs)

        for text in texts:
            for pattern_index, pattern in enumerate(_LEAK_PATTERNS):
                for match in pattern.findall(text):
                    issues["leak_tokens"].append(match)
                    if pattern_index == 0:
                        issues["unrendered_tags"].append(match)
        return issues

    def generate(self, context, output_path, include_toc=False, convert_to_pdf=False, required_fields=None):
        """
        Renders template with context, saves to output_path (.docx),
        optionally inserts TOC, and optionally renders PDF.
        Returns a dict with paths to generated files.
        """
        output_path = os.path.abspath(output_path)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # 1. Context validation
        validated_context, missing_fields = self.validate_context(context, required_fields=required_fields)

        # 2. Render Jinja2 Docx template
        doc_tpl = DocxTemplate(self.template_path)
        doc_tpl.render(validated_context)
        doc_tpl.save(output_path)

        # 3. Handle TOC insertion if requested
        if include_toc:
            doc = Document(output_path)
            if len(doc.paragraphs) > 0:
                new_para = doc.paragraphs[0].insert_paragraph_before('Table of Contents')
                new_para.style = 'Heading 1'
                toc_para = doc.paragraphs[1].insert_paragraph_before('')
                self._add_toc(toc_para)
                doc.paragraphs[2].insert_paragraph_before('').add_run().add_break()
                self._force_update_toc(doc)
                doc.save(output_path)

        # 4. Post-render validation -- on the FINAL saved file (after TOC),
        # not the pre-TOC intermediate, so it reflects what actually ships.
        issues = self._check_rendered_output(output_path)
        issues["missing_fields"] = missing_fields

        has_issues = bool(issues["structural_error"] or issues["leak_tokens"] or issues["missing_fields"])
        results = {
            "docx_path": output_path,
            "pdf_path": None,
            "status": "issues_found" if has_issues else "success",
            "issues": issues,
        }

        # 5. Optional PDF Conversion via LibreOffice
        if convert_to_pdf:
            pdf_path = self.convert_to_pdf(output_path)
            results["pdf_path"] = pdf_path

        return results

    def convert_to_pdf(self, docx_path):
        """Converts a .docx file to .pdf using LibreOffice soffice CLI."""
        docx_path = os.path.abspath(docx_path)
        output_dir = os.path.dirname(docx_path)
        base_name = os.path.splitext(os.path.basename(docx_path))[0]
        expected_pdf = os.path.join(output_dir, f"{base_name}.pdf")

        try:
            cmd = ["soffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", output_dir]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            if os.path.exists(expected_pdf):
                return expected_pdf
        except Exception as e:
            logger.warning(f"LibreOffice PDF conversion failed: {e}")

        return None
