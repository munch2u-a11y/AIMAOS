import os
import json
import re
import yaml
import docx
import logging
import threading
from datetime import datetime

logger = logging.getLogger(__name__)

CHUNK_PARAGRAPH_SIZE = 20  # Optimized for local 2B-8B LLMs (approx 300-500 tokens per chunk)

class TemplateReviewer:
    """
    Subagent for Alix-AI that performs background text reviews, template cleaning,
    and field schema optimization on Word .docx templates in small, token-bounded chunks.
    """
    def __init__(self, templates_dir="./templates", memory_dir="./workspace/.memory", llm_client=None):
        self.templates_dir = os.path.abspath(templates_dir)
        self.memory_dir = os.path.abspath(memory_dir)
        self.llm = llm_client
        self.notes_file = os.path.join(self.memory_dir, "template_notes.json")
        self.history_file = os.path.join(self.memory_dir, "template_reviews.json")
        
        os.makedirs(self.memory_dir, exist_ok=True)
        self.notes = self._load_json(self.notes_file, default=[])
        self.history = self._load_json(self.history_file, default=[])

    def _load_json(self, filepath, default):
        if os.path.exists(filepath):
            try:
                with open(filepath, "r") as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Error loading {filepath}: {e}")
        return default

    def _save_json(self, filepath, data):
        try:
            with open(filepath, "w") as f:
                json.dump(data, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving {filepath}: {e}")

    def add_review_note(self, template_name, note, priority="normal"):
        """Main agent or user calls this to queue a template for review/refinement."""
        entry = {
            "id": f"note_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            "timestamp": datetime.now().isoformat(),
            "template_name": template_name,
            "note": note,
            "priority": priority,
            "status": "pending"
        }
        self.notes.append(entry)
        self._save_json(self.notes_file, self.notes)
        logger.info(f"Added template review note for '{template_name}': {note}")
        return f"Queued review note for template '{template_name}'."

    def process_pending_notes(self):
        """Processes all pending review notes in small, token-optimized chunks."""
        pending = [n for n in self.notes if n.get("status") == "pending"]
        if not pending:
            return "No pending template review notes to process."

        results = []
        for note_entry in pending:
            t_name = note_entry["template_name"]
            res = self.review_and_refine_template(t_name, note=note_entry["note"])
            note_entry["status"] = "completed"
            note_entry["completed_at"] = datetime.now().isoformat()
            note_entry["result"] = res
            results.append(f"- Template [{t_name}]: {res}")

        self._save_json(self.notes_file, self.notes)
        return "\n".join(results)

    def review_and_refine_template(self, template_name, note=None):
        """
        Reviews a template chunk-by-chunk (20 paragraphs per chunk),
        detects unmapped fill-in blanks, updates template.yaml, and cleans footers/headers.
        """
        # Resolve template dir
        t_folder = os.path.join(self.templates_dir, template_name)
        docx_path = os.path.join(t_folder, "template.docx")
        yaml_path = os.path.join(t_folder, "template.yaml")

        if not os.path.exists(docx_path):
            # Try direct file match
            if os.path.exists(os.path.join(self.templates_dir, f"{template_name}.docx")):
                docx_path = os.path.join(self.templates_dir, f"{template_name}.docx")
                t_folder = os.path.dirname(docx_path)
            else:
                return f"Error: Template document for '{template_name}' not found."

        try:
            doc = docx.Document(docx_path)
            total_paragraphs = len(doc.paragraphs)
            modified_fields = set()
            unmapped_blanks = 0

            # Process document in small, token-optimized paragraph chunks
            for chunk_start in range(0, total_paragraphs, CHUNK_PARAGRAPH_SIZE):
                chunk_paras = doc.paragraphs[chunk_start : chunk_start + CHUNK_PARAGRAPH_SIZE]
                unmapped, fields_found = self._process_paragraph_chunk(chunk_paras, note=note)
                unmapped_blanks += unmapped
                modified_fields.update(fields_found)

            # Strip footers and headers if requested by note or default
            for s in doc.sections:
                for p in s.footer.paragraphs:
                    if "Instruction" in p.text or "Rules" in p.text or "Form 12." in p.text:
                        p.text = ""

            doc.save(docx_path)

            # Update or create template.yaml metadata
            yaml_content = {}
            if os.path.exists(yaml_path):
                try:
                    with open(yaml_path, "r") as f:
                        yaml_content = yaml.safe_load(f) or {}
                except Exception:
                    pass

            existing_fields = yaml_content.get("fields", {})
            if isinstance(existing_fields, list):
                existing_fields = {k: f"Value for {k}" for k in existing_fields}

            for f_name in modified_fields:
                if f_name not in existing_fields:
                    existing_fields[f_name] = f"Auto-detected variable for {f_name.replace('_', ' ')}"

            yaml_content["name"] = yaml_content.get("name", template_name.replace("_", " ").title())
            yaml_content["description"] = yaml_content.get("description", f"Refined template for {template_name}")
            yaml_content["fields"] = existing_fields
            yaml_content["last_reviewed_at"] = datetime.now().isoformat()

            os.makedirs(os.path.dirname(yaml_path), exist_ok=True)
            with open(yaml_path, "w") as f:
                yaml.dump(yaml_content, f, sort_keys=False)

            review_record = {
                "timestamp": datetime.now().isoformat(),
                "template_name": template_name,
                "paragraphs_analyzed": total_paragraphs,
                "unmapped_blanks_fixed": unmapped_blanks,
                "fields_active": list(existing_fields.keys()),
                "status": "success"
            }
            self.history.append(review_record)
            self._save_json(self.history_file, self.history)

            return f"Refined successfully. Analyzed {total_paragraphs} paragraphs in chunks of {CHUNK_PARAGRAPH_SIZE}. Fixed {unmapped_blanks} unmapped blanks. Active fields: {len(existing_fields)}."

        except Exception as e:
            return f"Error reviewing template '{template_name}': {e}"

    def _process_paragraph_chunk(self, paragraph_chunk, note=None):
        """Processes a single 20-paragraph chunk for local LLM token optimization."""
        unmapped_count = 0
        fields_found = set()

        for p in paragraph_chunk:
            text = p.text
            if not text.strip():
                continue

            # Detect existing Jinja2 variables
            for match in re.findall(r"\{\{\s*([a-zA-Z0-9_]+)", text):
                fields_found.add(match)

            # Heuristic cleaning of raw underline prompts in body text
            if "____" in text or "{" in text:
                if "Printed Name:" in text and "{{" not in text:
                    p.text = re.sub(r"_{3,}", " {{ client_name }}", p.text)
                    fields_found.add("client_name")
                    unmapped_count += 1
                elif "Address:" in text and "{{" not in text and "City" not in text:
                    p.text = re.sub(r"_{3,}", " {{ client_address }}", p.text)
                    fields_found.add("client_address")
                    unmapped_count += 1
                elif "Telephone Number:" in text and "{{" not in text:
                    p.text = re.sub(r"_{3,}", " {{ client_phone }}", p.text)
                    fields_found.add("client_phone")
                    unmapped_count += 1
                elif "E-Mail Address" in text and "{{" not in text:
                    p.text = re.sub(r"_{3,}", " {{ client_email }}", p.text)
                    fields_found.add("client_email")
                    unmapped_count += 1

        return unmapped_count, fields_found

    def run_background_worker(self):
        """Background thread loop for asynchronous template reviews."""
        t = threading.Thread(target=self.process_pending_notes, daemon=True)
        t.start()
        return t
