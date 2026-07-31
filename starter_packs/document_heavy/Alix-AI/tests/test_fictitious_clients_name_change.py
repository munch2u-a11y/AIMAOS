import os

def _find_aimaos_root():
    p = os.path.dirname(os.path.abspath(__file__))
    while p != os.path.dirname(p) and not os.path.exists(os.path.join(p, "aimaos_config.yaml")):
        p = os.path.dirname(p)
    return p
AIMAOS_ROOT = os.environ.get("AIMAOS_ROOT") or _find_aimaos_root()
import sys
import docx

sys.path.insert(0, AIMAOS_ROOT)
from core.tools import ToolRegistry

def run_fictitious_client_tests():
    registry = ToolRegistry(os.path.join(AIMAOS_ROOT, "Alix-AI/tools"))

    print("====================================================================")
    print("ALIX-AI MULTI-COUNTY FICTITIOUS CLIENT NAME CHANGE TEST SUITE")
    print("====================================================================")

    clients = [
        {
            "intake_file": os.path.join(AIMAOS_ROOT, "Alix-AI/workspace/inbox/intake_alex_sample_leon.txt"),
            "client_name": "Alex Sample",
            "new_name": "Alex Newname",
            "county": "Leon",
            "circuit_number": "2nd",
            "case_number": "2026-DR-002101",
            "division": "Family Division",
            "client_address": "1420 Timberlane Road, Tallahassee, FL 32312",
            "client_phone": "(850) 555-0144",
            "date_of_birth": "March 12, 1991",
            "output_slug": "alex_sample_leon"
        },
        {
            "intake_file": os.path.join(AIMAOS_ROOT, "Alix-AI/workspace/inbox/intake_dana_sample_duval.txt"),
            "client_name": "Dana Sample",
            "new_name": "Elena Vance",
            "county": "Duval",
            "circuit_number": "4th",
            "case_number": "2026-DR-005432",
            "division": "Family Division",
            "client_address": "850 Ocean Boulevard, Jacksonville, FL 32216",
            "client_phone": "(904) 555-0188",
            "date_of_birth": "July 22, 1994",
            "output_slug": "dana_sample_duval"
        },
        {
            "intake_file": os.path.join(AIMAOS_ROOT, "Alix-AI/workspace/inbox/intake_morgan_sample_st_johns.txt"),
            "client_name": "Morgan Sample",
            "new_name": "Morgan Sample-Blackwood",
            "county": "St. Johns",
            "circuit_number": "7th",
            "case_number": "2026-DR-008910",
            "division": "Family Division",
            "client_address": "310 Avenida Menendez, St. Augustine, FL 32084",
            "client_phone": "(904) 555-0199",
            "date_of_birth": "November 05, 1988",
            "output_slug": "morgan_sample_st_johns"
        }
    ]

    results_summary = []

    for idx, c in enumerate(clients, 1):
        print(f"\n--- CLIENT {idx}: {c['client_name']} ({c['county']} County, {c['circuit_number']} Judicial Circuit) ---")

        # Step 1: Ingest Intake Form into mRAG Layer 1
        ingest_res = registry.execute_tool("ingest_document", {
            "file_path": c["intake_file"],
            "client_name": c["client_name"]
        })
        print(f"[mRAG INGESTION]: {ingest_res}")

        # Step 2: Populate Form 12.982(a) Petition for Change of Name (Adult)
        petition_res = registry.execute_tool("populate_template", {
            "template_name": "form_12_982_a",
            "context": {
                "client_name": c["client_name"],
                "new_name": c["new_name"],
                "county": c["county"],
                "circuit_number": c["circuit_number"],
                "case_number": c["case_number"],
                "division": c["division"],
                "client_address": c["client_address"],
                "client_phone": c["client_phone"],
                "date_of_birth": c["date_of_birth"]
            },
            "output_name": f"{c['output_slug']}_petition_name_change"
        })
        print(f"[POPULATE PETITION]:\n{petition_res}")

        # Extract petition docx path
        petition_docx = None
        for line in petition_res.split("\n"):
            if "Word Document:" in line:
                petition_docx = line.split("Word Document:")[1].strip()

        # Step 3: Populate Form 12.982(b) Final Judgment of Change of Name (Adult)
        judgment_res = registry.execute_tool("populate_template", {
            "template_name": "form_12_982_b",
            "context": {
                "client_name": c["client_name"],
                "new_name": c["new_name"],
                "county": c["county"],
                "circuit_number": c["circuit_number"],
                "case_number": c["case_number"],
                "division": c["division"]
            },
            "output_name": f"{c['output_slug']}_final_judgment"
        })
        print(f"[POPULATE FINAL JUDGMENT]:\n{judgment_res}")

        # Extract judgment docx path
        judgment_docx = None
        for line in judgment_res.split("\n"):
            if "Word Document:" in line:
                judgment_docx = line.split("Word Document:")[1].strip()

        # Step 4: Dispatch both documents to client archive directory
        if petition_docx:
            d_res1 = registry.execute_tool("dispatch_document", {
                "file_path": petition_docx,
                "client_name": c["client_name"],
                "notes": f"Petition for Adult Name Change for {c['client_name']} ({c['county']} County)"
            })
            print(f"[DISPATCH PETITION]: {d_res1.splitlines()[0]}")

        if judgment_docx:
            d_res2 = registry.execute_tool("dispatch_document", {
                "file_path": judgment_docx,
                "client_name": c["client_name"],
                "notes": f"Final Judgment of Adult Name Change for {c['client_name']} ({c['county']} County)"
            })
            print(f"[DISPATCH JUDGMENT]: {d_res2.splitlines()[0]}")

        results_summary.append({
            "client": c["client_name"],
            "county": c["county"],
            "circuit": c["circuit_number"],
            "petition_docx": petition_docx,
            "judgment_docx": judgment_docx
        })

    # Step 5: Verification Audit across generated documents
    print("\n====================================================================")
    print("VERIFICATION AUDIT OF GENERATED COURT MOTIONS & JUDGMENTS")
    print("====================================================================")

    for item in results_summary:
        print(f"\nAudit for Client: {item['client']} ({item['county']} County, {item['circuit']} Judicial Circuit)")
        for doc_type, path in [("Petition", item["petition_docx"]), ("Final Judgment", item["judgment_docx"])]:
            if not path or not os.path.exists(path):
                print(f"  ❌ {doc_type}: Missing file ({path})")
                continue

            doc = docx.Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)

            court_header_ok = "IN THE CIRCUIT COURT" in full_text
            circuit_ok = item["circuit"] in full_text
            county_ok = item["county"] in full_text
            instructions_removed = "When should this form be used?" not in full_text

            footer_clean = True
            for s in doc.sections:
                for p in s.footer.paragraphs:
                    if p.text.strip():
                        footer_clean = False

            print(f"  • {doc_type} Docx: {os.path.basename(path)}")
            print(f"    - Court Header Present: {court_header_ok}")
            print(f"    - Correct Circuit ({item['circuit']}): {circuit_ok}")
            print(f"    - Correct County ({item['county']}): {county_ok}")
            print(f"    - Instructions Stripped: {instructions_removed}")
            print(f"    - Footers Clean: {footer_clean}")

if __name__ == "__main__":
    run_fictitious_client_tests()
