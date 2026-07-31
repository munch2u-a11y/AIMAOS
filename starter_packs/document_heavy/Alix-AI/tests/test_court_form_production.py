import os

def _find_aimaos_root():
    p = os.path.dirname(os.path.abspath(__file__))
    while p != os.path.dirname(p) and not os.path.exists(os.path.join(p, "aimaos_config.yaml")):
        p = os.path.dirname(p)
    return p
AIMAOS_ROOT = os.environ.get("AIMAOS_ROOT") or _find_aimaos_root()
import sys
import docx

sys.path.insert(0, AIMAOS_ROOT)
from core.tools import ToolRegistry

def test_court_forms():
    registry = ToolRegistry(os.path.join(AIMAOS_ROOT, "Alix-AI/tools"))

    print("=== TESTING COURT FORM PRODUCTION ON CONVERTED TEMPLATES ===")

    # Test 1: Adult Name Change Petition (Form 12.982a)
    print("\n1. Testing Form 12.982(a) - Petition for Change of Name (Adult)...")
    res1 = registry.execute_tool("populate_template", {
        "template_name": "form_12_982_a",
        "context": {
            "client_name": "Lee Sample",
            "county": "Volusia",
            "circuit_number": "7th",
            "case_number": "2026-DR-004122",
            "division": "Family Law",
            "client_address": "789 Stability Street, Daytona Beach, FL 32114",
            "client_phone": "555-888-7777",
            "date_of_birth": "October 14, 1985"
        },
        "output_name": "lee_sample_name_change_petition"
    })
    print(res1)

    # Test 2: Notice of Related Cases (Form 12.900h)
    print("\n2. Testing Form 12.900(h) - Notice of Related Cases...")
    res2 = registry.execute_tool("populate_template", {
        "template_name": "form_12_900_h",
        "context": {
            "client_name": "Jane Doe",
            "petitioner_name": "Jane Doe",
            "respondent_name": "John Doe",
            "county": "Volusia",
            "circuit_number": "7th",
            "case_number": "2026-DR-009811",
            "division": "Family Law"
        },
        "output_name": "jane_doe_notice_of_related_cases"
    })
    print(res2)

    # Test 3: Joint Petition for Simplified Dissolution of Marriage (Form 12.901a)
    print("\n3. Testing Form 12.901(a) - Joint Petition for Simplified Dissolution...")
    res3 = registry.execute_tool("populate_template", {
        "template_name": "form_12_901_a",
        "context": {
            "client_name": "John Smith & Sarah Smith",
            "husband_name": "John Smith",
            "wife_name": "Sarah Smith",
            "county": "Volusia",
            "circuit_number": "7th",
            "case_number": "2026-DR-001290",
            "division": "Family Law"
        },
        "output_name": "smith_simplified_dissolution_petition"
    })
    print(res3)

    # Verification checks on output docx files
    print("\n=== RUNNING VERIFICATION AUDIT ON GENERATED DOCUMENTS ===")
    test_outputs = [
        os.path.join(AIMAOS_ROOT, "Alix-AI/workspace/output/lee_sample_name_change_petition.docx"),
        os.path.join(AIMAOS_ROOT, "Alix-AI/workspace/output/jane_doe_notice_of_related_cases.docx"),
        os.path.join(AIMAOS_ROOT, "Alix-AI/workspace/output/smith_simplified_dissolution_petition.docx")
    ]

    for path in test_outputs:
        if not os.path.exists(path):
            print(f"FAILED: Output file missing: {path}")
            continue

        doc = docx.Document(path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

        # Ensure instruction section is gone
        has_instruction = "When should this form be used?" in full_text
        has_court_header = "IN THE CIRCUIT COURT" in full_text

        # Check footers
        footer_text = ""
        for s in doc.sections:
            for p in s.footer.paragraphs:
                footer_text += p.text.strip()

        print(f"\nAudit for: {os.path.basename(path)}")
        print(f"  • Total Paragraphs: {len(doc.paragraphs)}")
        print(f"  • Starts with Court Header ('IN THE CIRCUIT COURT'): {has_court_header}")
        print(f"  • Instructions Removed ('When should this form be used?'): {not has_instruction}")
        print(f"  • Instruction Footers Cleaned: {footer_text == ''}")
        print(f"  • First paragraph: {doc.paragraphs[0].text[:80] if doc.paragraphs else 'Empty'}")

if __name__ == "__main__":
    test_court_forms()
