"""Bounded, local-only validation and text extraction for intake files."""
from __future__ import annotations

import csv
import io
import os
import zipfile
from dataclasses import dataclass

from core.security import SecurityValidationError

MAX_EXTRACTED_CHARS = 100_000
MAX_ARCHIVE_ENTRIES = 5_000
MAX_ARCHIVE_EXPANDED_BYTES = 100 * 1024 * 1024
MAX_COMPRESSION_RATIO = 1_000


@dataclass
class ExtractionResult:
    text: str
    status: str
    detail: str


def _validate_office_archive(raw: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_ARCHIVE_ENTRIES:
                raise SecurityValidationError("Office document contains too many archive entries.")
            expanded = sum(entry.file_size for entry in entries)
            compressed = max(1, sum(entry.compress_size for entry in entries))
            if expanded > MAX_ARCHIVE_EXPANDED_BYTES or expanded / compressed > MAX_COMPRESSION_RATIO:
                raise SecurityValidationError("Office document expands beyond the safe processing limit.")
            for entry in entries:
                normalized = entry.filename.replace("\\", "/")
                if normalized.startswith("/") or "../" in f"/{normalized}":
                    raise SecurityValidationError("Office document contains an unsafe archive path.")
    except zipfile.BadZipFile as exc:
        raise SecurityValidationError("Office document is not a valid DOCX/XLSX archive.") from exc


def validate_upload_content(filename: str, raw: bytes) -> None:
    """Reject common extension/signature mismatches before data reaches parsers."""
    extension = os.path.splitext(filename)[1].lower()
    prefix = raw[:16]
    text_extensions = {".txt", ".md", ".csv", ".json", ".yaml", ".yml"}
    if extension in text_extensions:
        if b"\x00" in raw[:64_000]:
            raise SecurityValidationError("Text upload contains binary data.")
        return
    if extension == ".rtf" and not raw.lstrip().startswith(b"{\\rtf"):
        raise SecurityValidationError("File content does not match the .rtf extension.")
    if extension == ".pdf" and not raw.lstrip().startswith(b"%PDF-"):
        raise SecurityValidationError("File content does not match the .pdf extension.")
    if extension in {".docx", ".xlsx"}:
        if not zipfile.is_zipfile(io.BytesIO(raw)):
            raise SecurityValidationError(f"File content does not match the {extension} extension.")
        _validate_office_archive(raw)
    if extension in {".doc", ".xls"} and not prefix.startswith(bytes.fromhex("D0CF11E0A1B11AE1")):
        raise SecurityValidationError(f"File content does not match the {extension} extension.")
    signatures = {
        ".png": (b"\x89PNG\r\n\x1a\n",),
        ".jpg": (b"\xff\xd8\xff",),
        ".jpeg": (b"\xff\xd8\xff",),
        ".tif": (b"II*\x00", b"MM\x00*"),
        ".tiff": (b"II*\x00", b"MM\x00*"),
        ".wav": (b"RIFF",),
        ".webm": (b"\x1aE\xdf\xa3",),
    }
    if extension in signatures and not any(prefix.startswith(signature) for signature in signatures[extension]):
        raise SecurityValidationError(f"File content does not match the {extension} extension.")
    if extension == ".mp3" and not (prefix.startswith(b"ID3") or prefix.startswith(b"\xff")):
        raise SecurityValidationError("File content does not match the .mp3 extension.")
    if extension == ".m4a" and b"ftyp" not in raw[4:16]:
        raise SecurityValidationError("File content does not match the .m4a extension.")


def extract_document_text(path: str) -> ExtractionResult:
    extension = os.path.splitext(path)[1].lower()
    if extension in {".txt", ".md", ".json", ".yaml", ".yml", ".rtf"}:
        with open(path, "r", encoding="utf-8", errors="replace") as handle:
            text = handle.read(MAX_EXTRACTED_CHARS + 1)
        truncated = len(text) > MAX_EXTRACTED_CHARS
        return ExtractionResult(
            text=text[:MAX_EXTRACTED_CHARS], status="extracted",
            detail="Text extracted locally" + (" and truncated to the review limit" if truncated else ""),
        )
    if extension == ".csv":
        rows = []
        with open(path, "r", encoding="utf-8", errors="replace", newline="") as handle:
            for row in csv.reader(handle):
                rows.append(" | ".join(row))
                if sum(len(item) for item in rows) >= MAX_EXTRACTED_CHARS:
                    break
        return ExtractionResult("\n".join(rows)[:MAX_EXTRACTED_CHARS], "extracted", "CSV text extracted locally")
    if extension == ".docx":
        import docx
        document = docx.Document(path)
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return ExtractionResult("\n".join(parts)[:MAX_EXTRACTED_CHARS], "extracted", "DOCX text extracted locally")
    if extension == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(path)
        parts = []
        for page in reader.pages[:50]:
            parts.append(page.extract_text() or "")
            if sum(len(item) for item in parts) >= MAX_EXTRACTED_CHARS:
                break
        text = "\n".join(parts)[:MAX_EXTRACTED_CHARS]
        if not text.strip():
            return ExtractionResult("", "manual_review_required", "PDF has no extractable text layer")
        return ExtractionResult(text, "extracted", "PDF text layer extracted locally")
    return ExtractionResult(
        "", "manual_review_required",
        f"Automated text extraction is not enabled for {extension or 'this file type'}",
    )
